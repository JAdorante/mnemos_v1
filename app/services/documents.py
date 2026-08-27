"""Read-my-documents ingestion — turn the TEXT of a user's files into memory.

The app/git/bookmark scan (services/onboarding_scan.py) reads only *metadata*:
which apps exist, which folders are code projects, the git identity. This module
crosses the line the scan deliberately won't — it reads the actual *content* of
documents (PDF / Word / plain notes) and runs it through the SAME extraction
pipeline that mines speech, so a resume, a spec, or a meeting-notes doc becomes
structured tasks / commitments / claims / entities you can ask about.

Because it reads content, it is the most sensitive enrichment source, so:
  * it is EXPLICIT opt-in — never part of the automatic scan, always its own
    consented action (a wizard checkbox + the /onboarding/documents endpoint);
  * document text is sent to the extraction model — the same trust boundary as
    every transcript, disclosed at the point of consent;
  * every emitted fact lands UNREVIEWED (epistemic=extracted), so it shows up in
    the Memory Console's approve/edit/dismiss loop rather than as settled truth;
  * everything is reversible — all events carry source='documents.scan'.

General-purpose by construction (the generality rule): the scanner reads
whatever machine it runs on — the OS document folders, whatever files are there
— with zero user-specific logic. On another person's machine it learns their
documents. All limits live in DocumentsConfig (env), not in code.

Every function is best-effort and never raises: a file it can't parse is skipped,
and a scan that finds nothing just reports zero.
"""
from __future__ import annotations

import hashlib
import json
import os
import re
import time
from pathlib import Path
from typing import Iterable

from app.config import settings

SOURCE = "documents.scan"

# Folder names never worth reading — dependency trees, build junk, caches, VCS
# internals. Mirrors onboarding_scan._PROJECT_SKIP plus doc-specific noise.
_SKIP_DIRS = frozenset({
    "node_modules", "venv", ".venv", "env", ".git", "__pycache__", "dist",
    "build", "target", ".idea", ".vscode", "site-packages", ".cache", ".next",
    ".gradle", ".m2", "appdata", "temp", "tmp", "$recycle.bin", ".ipynb_checkpoints",
})

# The app's own repo tree — never ingest our own source/docs (on a developer's
# machine the README / architecture notes / requirements sit right here, and
# reading them back in floods the graph with "FastAPI server", "data/quill.db",
# and friends). parents[2] = <repo>/app/services/documents.py -> <repo>.
try:
    _PROJECT_ROOT = Path(__file__).resolve().parents[2]
except Exception:
    _PROJECT_ROOT = None

# Developer / build / dependency / log files that are documents by extension but
# not the USER's documents. Skipped wherever they sit (loose copies in Downloads
# count too), so "read my documents" targets notes/essays/work docs, not code
# artifacts. General knowledge, like _SKIP_DIRS.
_DEV_FILE_RE = re.compile(
    r"""(?ix) ^ (
        readme | changelog | change[-_]?log | history | license | licence |
        contributing | authors | notice | code[-_]?of[-_]?conduct | manifest |
        requirements | pipfile | setup | pyproject | tox | makefile | dockerfile |
        tsconfig | package(-lock)? | pnpm-lock | yarn | \.env | \.gitignore |
        skill
    ) ($ | [-_. ]) |
    ( \.lock$ | [-_.]lock[-_.] | \.log$ | install[-_.]?log | package-lock\.json$ )
    """,
    re.X | re.I)


def _is_dev_file(path: Path) -> bool:
    """A code/build/dependency/log file — not a user document."""
    if _PROJECT_ROOT is not None and _within(_PROJECT_ROOT, path):
        return True
    return bool(_DEV_FILE_RE.search(path.name))


# --- roots ------------------------------------------------------------------
def _default_roots() -> list[Path]:
    """The OS's usual document folders. Read from the environment, so this is the
    same on any Windows machine and degrades gracefully elsewhere."""
    home = Path(os.path.expanduser("~"))
    candidates = [home / "Documents", home / "Desktop", home / "Downloads"]
    # OneDrive redirects Documents/Desktop on many Windows setups.
    od = os.environ.get("OneDrive") or os.environ.get("OneDriveConsumer")
    if od:
        candidates += [Path(od) / "Documents", Path(od) / "Desktop"]
    out, seen = [], set()
    for p in candidates:
        try:
            rp = p.resolve()
        except Exception:
            continue
        if rp.is_dir() and str(rp).lower() not in seen:
            seen.add(str(rp).lower())
            out.append(rp)
    return out


def _configured_roots() -> list[Path]:
    raw = settings.documents.roots_raw.strip()
    if not raw:
        return _default_roots()
    out = []
    for part in raw.split(";"):
        part = part.strip()
        if not part:
            continue
        try:
            p = Path(part).resolve()
        except Exception:
            continue
        if p.is_dir():
            out.append(p)
    return out or _default_roots()


def _within(root: Path, child: Path) -> bool:
    """Is `child` inside `root`? Guards against symlink/`..` escapes so ingestion
    can never wander outside a document root it was pointed at."""
    try:
        child.resolve().relative_to(root.resolve())
        return True
    except Exception:
        return False


# --- discovery --------------------------------------------------------------
def discover(roots: list[Path] | None = None, exts: Iterable[str] | None = None,
             max_docs: int | None = None, max_bytes: int | None = None,
             max_depth: int | None = None) -> list[Path]:
    """Find candidate document files under the roots: matching extension, under
    the size cap, not in a junk dir. Newest-modified first (recent work is the
    most useful signal), then capped. Pure filesystem read — no ingestion."""
    cfg = settings.documents
    roots = roots if roots is not None else _configured_roots()
    exts = frozenset(e.lower() for e in exts) if exts is not None else cfg.exts
    max_docs = cfg.max_docs if max_docs is None else max_docs
    max_bytes = cfg.max_bytes if max_bytes is None else max_bytes
    max_depth = cfg.max_depth if max_depth is None else max_depth

    found: list[tuple[float, Path]] = []
    for root in roots:
        root = root.resolve()
        base_depth = len(root.parts)
        for dirpath, dirnames, filenames in os.walk(root):
            d = Path(dirpath)
            # Prune junk + hidden dirs in place so os.walk doesn't descend them.
            dirnames[:] = [n for n in dirnames
                           if n.lower() not in _SKIP_DIRS and not n.startswith(".")]
            if len(d.parts) - base_depth >= max_depth:
                dirnames[:] = []
            for fn in filenames:
                if Path(fn).suffix.lower() not in exts:
                    continue
                fp = d / fn
                if _is_dev_file(fp):     # code/build/dep/log/our-own-repo — not a user doc
                    continue
                try:
                    st = fp.stat()
                except OSError:
                    continue
                if st.st_size <= 0 or st.st_size > max_bytes:
                    continue
                if not _within(root, fp):
                    continue
                found.append((st.st_mtime, fp))
    found.sort(key=lambda t: t[0], reverse=True)
    # Dedup by resolved path (roots can overlap, e.g. Desktop under OneDrive).
    out, seen = [], set()
    for _, fp in found:
        key = str(fp).lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(fp)
        if len(out) >= max_docs:
            break
    return out


# --- text extraction --------------------------------------------------------
def _read_plain(path: Path, cap: int) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")[:cap]
    except Exception:
        try:
            return path.read_text(encoding="latin-1", errors="ignore")[:cap]
        except Exception:
            return ""


def _page_text(page) -> str:
    """One page's text, preferring pypdf's layout mode.

    Word-processor exports (Google Docs, Pages) place each word as its own
    positioned text object, and pypdf's default mode emits a newline after
    every one — "AI\\n \\nin\\n \\nthe\\n \\ncompany." Readable prose becomes a
    column of single words, which burns a third of the char cap on whitespace
    and is what an LLM then has to summarize. Layout mode reconstructs the
    lines. Falls back to the default mode when layout is unavailable (older
    pypdf) or throws on a malformed page.
    """
    try:
        t = page.extract_text(extraction_mode="layout") or ""
    except Exception:
        t = ""
    if t.strip():
        return t
    try:
        return page.extract_text() or ""
    except Exception:
        return ""


def _read_pdf(path: Path, cap: int) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        return ""
    try:
        reader = PdfReader(str(path))
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")  # only unlock no-password encryption
            except Exception:
                return ""
        parts, total = [], 0
        for page in reader.pages:
            t = _page_text(page)
            if not t:
                continue
            parts.append(t)
            total += len(t)
            if total >= cap:
                break
        return "\n".join(parts)[:cap]
    except Exception:
        return ""


def _read_docx(path: Path, cap: int) -> str:
    try:
        import docx  # python-docx
    except Exception:
        return ""
    try:
        doc = docx.Document(str(path))
        parts, total = [], 0
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if not t:
                continue
            parts.append(t)
            total += len(t)
            if total >= cap:
                break
        # Table cells often hold the real content (invoices, specs).
        if total < cap:
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        line = " | ".join(cells)
                        parts.append(line)
                        total += len(line)
                        if total >= cap:
                            break
                if total >= cap:
                    break
        return "\n".join(parts)[:cap]
    except Exception:
        return ""


_READERS = {
    ".txt": _read_plain, ".md": _read_plain, ".markdown": _read_plain,
    ".rst": _read_plain, ".text": _read_plain, ".log": _read_plain,
    ".pdf": _read_pdf, ".docx": _read_docx,
}


def extract_text(path: Path, cap: int | None = None) -> str:
    """Best-effort plain text for one document, bounded to `cap` chars. Returns
    '' for an unsupported type, an unreadable/corrupt file, or an image-only PDF
    (no embedded text — OCR is out of scope here)."""
    cap = settings.documents.max_chars if cap is None else cap
    reader = _READERS.get(path.suffix.lower())
    if reader is None:
        return ""
    return (reader(path, cap) or "").strip()


# --- chunking ---------------------------------------------------------------
def chunk_text(text: str, size: int | None = None) -> list[str]:
    """Split into ~`size`-char chunks on paragraph boundaries, so each extractor
    call sees a coherent block. Falls back to a hard slice for a giant unbroken
    paragraph. Empty/blank chunks are dropped."""
    size = settings.documents.chunk_chars if size is None else size
    text = (text or "").strip()
    if not text:
        return []
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paras:
        paras = [text]
    chunks: list[str] = []
    cur = ""
    for p in paras:
        if len(p) > size:
            if cur:
                chunks.append(cur)
                cur = ""
            for i in range(0, len(p), size):
                chunks.append(p[i:i + size])
            continue
        if cur and len(cur) + len(p) + 2 > size:
            chunks.append(cur)
            cur = p
        else:
            cur = f"{cur}\n\n{p}" if cur else p
    if cur:
        chunks.append(cur)
    return [c for c in chunks if c.strip()]


# --- idempotency ledger -----------------------------------------------------
def _load_state() -> dict:
    p = Path(settings.documents.state_path)
    try:
        return json.loads(p.read_text(encoding="utf-8")) if p.is_file() else {}
    except Exception:
        return {}


def _save_state(state: dict) -> None:
    p = Path(settings.documents.state_path)
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(json.dumps(state, indent=2), encoding="utf-8")
    except Exception as exc:
        print(f"[documents] state save skipped ({exc}).")


def _file_key(path: Path) -> str:
    return hashlib.sha1(str(path.resolve()).lower().encode("utf-8")).hexdigest()


def _sig(path: Path) -> str:
    """A content signature (mtime + size) so an unchanged file is skipped but an
    edited one re-ingests on the next run."""
    try:
        st = path.stat()
        return f"{int(st.st_mtime)}:{st.st_size}"
    except OSError:
        return ""


# --- fact persistence (reuses the audio extractor, minus the live-offer path) --
def _chat_people_names(text: str) -> list[str]:
    """Multi-word proper names in typed chat worth minting as people.

    The extractor often parks a meeting as a commitment/claim with empty or
    self/unknown parties and never passes the counterparty into
    resolve_person_mention — so "meeting with Andy Karos" never becomes a
    People row. Harvesting multi-word names from the raw chat text closes that
    gap without minting every capitalized single token ("Remember", "Today").
    """
    from app.services.answer_check import extract_name_tokens
    from app.services import name_quality as nq
    from app.services import self_profile

    out: list[str] = []
    seen: set[str] = set()
    for tok in extract_name_tokens(text or ""):
        if " " not in (tok or ""):
            continue
        if not nq.is_plausible_person(tok):
            continue
        if self_profile.is_self_name(tok):
            continue
        key = tok.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(tok)
    return out


def _persist_facts(store, facts: dict, anchor: int, chunk: str, now: float,
                   index: bool = True, *, event_source: str = "document",
                   window: str = "") -> int:
    """Write the extractor's output for one chunk: tasks / commitments / claims +
    the entity/relation graph, each anchored to the document event for provenance.

    Deliberately does NOT call task_offer (that's for tasks heard *just now* in
    conversation — a bulk document import must not fire a burst of 'run this?'
    chat offers). Otherwise it mirrors Extractor._persist, including the
    write-time hygiene gate (confidence floor / span faithfulness / dedup /
    supersede — see services/fact_gate.py). `index` gates writing into the
    shared vector store — off when a caller passes a throwaway store (its fact
    ids wouldn't line up with the live index), like enrich() does; the
    dedup/supersede probe is skipped too then, for the same id-mismatch reason.

    People v2: `event_source` / `window` feed source_policy (screen vs docs vs
    news) so person minting and contact attribution stay policy-gated."""
    from app.services.extractor import _coerce_due, _index_fact, extractor
    from app.services.resolution import resolver

    def _person(name: str, *, role: str = "unknown", boost: float = 0.55):
        name = (name or "").strip()
        if not name:
            return None
        from app.services import self_profile
        from app.services.people_pipeline import enabled, resolve_person_mention
        if self_profile.is_self_name(name):
            return self_profile.self_person_id(store)
        if enabled():
            res = resolve_person_mention(
                name, store=store, event_id=anchor,
                event_source=event_source, window=window, text=chunk,
                grammatical_role=role, now=now,
                relationship_boost=boost)
            if res.person_id and chunk:
                from app.services.people_pipeline import attribute_contacts_from_text
                attribute_contacts_from_text(
                    chunk, store=store, person_id=res.person_id,
                    person_name=name, event_id=anchor, now=now,
                    event_source=event_source, window=window)
            return res.person_id
        return resolver.resolve_person(name, ts=now)

    def _fill_commitment_parties(fid: int, *, from_pid, to_pid) -> None:
        """Backfill NULL party columns on a deduped commitment (never overwrite)."""
        if not fid or (from_pid is None and to_pid is None):
            return
        try:
            with store._lock:
                if from_pid is not None:
                    store._conn.execute(
                        "UPDATE commitments SET from_person_id=? "
                        "WHERE fact_id=? AND from_person_id IS NULL",
                        (int(from_pid), int(fid)))
                if to_pid is not None:
                    store._conn.execute(
                        "UPDATE commitments SET to_person_id=? "
                        "WHERE fact_id=? AND to_person_id IS NULL",
                        (int(to_pid), int(fid)))
                store._conn.commit()
            for pid in (from_pid, to_pid):
                if pid:
                    store.mark_graph_dirty("person", int(pid), ts=now)
            store.mark_graph_dirty("fact", int(fid), ts=now)
        except Exception as exc:
            print(f"[documents] commitment party backfill skipped ({exc}).")

    def _fill_task_owner(fid: int, owner_pid) -> None:
        if not fid or owner_pid is None:
            return
        try:
            with store._lock:
                store._conn.execute(
                    "UPDATE tasks SET owner_person_id=? "
                    "WHERE fact_id=? AND owner_person_id IS NULL",
                    (int(owner_pid), int(fid)))
                store._conn.commit()
            store.mark_graph_dirty("person", int(owner_pid), ts=now)
            store.mark_graph_dirty("fact", int(fid), ts=now)
        except Exception as exc:
            print(f"[documents] task owner backfill skipped ({exc}).")

    def _idx(fid: int, kind: str, text: str) -> None:
        if index:
            _index_fact(store, fid, kind, text, now)

    def _gate(kind: str, item: dict):
        try:
            from app.config import settings
            from app.services.fact_gate import Verdict, gate_fact
            if index:
                return gate_fact(kind, item.get("text") or "",
                                 item.get("confidence"),
                                 item.get("source_span", ""), chunk)
            # Throwaway store: confidence floor only — no probe, no span gate
            # (enrich() previews shouldn't drop rows the live pass would keep).
            conf = item.get("confidence")
            cfg = settings.facts
            if cfg.min_conf > 0 and conf is not None and conf < cfg.min_conf:
                return Verdict("drop", "confidence below floor")
            return Verdict("insert")
        except Exception:
            class _Insert:  # duck-typed insert verdict
                action = "insert"
                dup_fact_id = None
                supersede_ids: tuple = ()
            return _Insert()

    def _apply(v, fid: int) -> None:
        for old in getattr(v, "supersede_ids", ()) or ():
            store.supersede_fact(old, fid, now)

    n = 0
    for t in facts.get("tasks", []):
        if not t.get("text"):
            continue
        v = _gate("task", t)
        if v.action == "drop":
            continue
        if v.action == "dedup":
            store.touch_fact(v.dup_fact_id, now, t.get("confidence"))
            # Still mint/link the owner — dedupe must not skip people.
            _fill_task_owner(
                v.dup_fact_id,
                _person(t.get("owner", ""), role="owner", boost=0.85))
            continue
        # Boosts match the live speech extractor: people minting requires
        # relationship_boost >= 0.70 (CREATE_RELEVANCE). The old 0.65/0.6
        # floors could only leave_open even when source_policy allowed minting.
        fid = store.add_task(
            t["text"], source_event_id=anchor, source_span=t.get("source_span", ""),
            confidence=t.get("confidence"),
            owner_person_id=_person(t.get("owner", ""), role="owner", boost=0.85),
            due=_coerce_due(t.get("due")), extracted_at=now)
        _apply(v, fid)
        _idx(fid, "task", t["text"])
        extractor._record_faithfulness(t, chunk)
        n += 1
    for c in facts.get("commitments", []):
        if not c.get("text"):
            continue
        v = _gate("commitment", c)
        if v.action == "drop":
            continue
        if v.action == "dedup":
            store.touch_fact(v.dup_fact_id, now, c.get("confidence"))
            _fill_commitment_parties(
                v.dup_fact_id,
                from_pid=_person(c.get("from_person", ""), role="from", boost=0.8),
                to_pid=_person(c.get("to_person", ""), role="to", boost=0.75))
            continue
        fid = store.add_commitment(
            c["text"], source_event_id=anchor, source_span=c.get("source_span", ""),
            confidence=c.get("confidence"),
            from_person_id=_person(c.get("from_person", ""), role="from", boost=0.8),
            to_person_id=_person(c.get("to_person", ""), role="to", boost=0.75),
            due=_coerce_due(c.get("due")),
            extracted_at=now)
        _apply(v, fid)
        _idx(fid, "commitment", c["text"])
        extractor._record_faithfulness(c, chunk)
        n += 1
    for cl in facts.get("claims", []):
        if not cl.get("text"):
            continue
        v = _gate("claim", cl)
        if v.action == "drop":
            continue
        if v.action == "dedup":
            store.touch_fact(v.dup_fact_id, now, cl.get("confidence"))
            continue
        fid = store.add_claim(
            cl["text"], source_event_id=anchor, source_span=cl.get("source_span", ""),
            confidence=cl.get("confidence"), extracted_at=now)
        _apply(v, fid)
        _idx(fid, "claim", cl["text"])
        extractor._record_faithfulness(cl, chunk)
        # First-person claim (typed chat / user's own documents) → self node.
        # `index` doubles as the live-store signal: never link on a throwaway.
        if index:
            from app.services import self_profile
            if self_profile.is_first_person(cl["text"]):
                self_profile.link_self(store, fid, now)
        n += 1
    # Entity nodes + asserted relation edges (the graph's non-person side).
    # Pass source_policy context so relation persons obey news/feed gates.
    try:
        extractor._persist_entities(
            facts, anchor, now,
            event_source=event_source, window=window, text=chunk)
    except Exception as exc:
        print(f"[documents] entity persist skipped ({exc}).")
    # Typed chat: mint multi-word names from the raw turn even when the
    # extractor omitted them as from/to (common for "meeting with X").
    if "chat" in (event_source or "").lower():
        for name in _chat_people_names(chunk):
            if _person(name, role="relation", boost=0.75):
                n += 1
    return n


# --- public API -------------------------------------------------------------
def roots() -> list[Path]:
    """The document folders an ingest will read (configured, else OS defaults).
    Public so the availability probe can show the user which folders are in scope
    before they consent."""
    return _configured_roots()


def preview(roots: list[Path] | None = None, limit: int | None = None) -> list[dict]:
    """Read-only listing of what an ingest WOULD read: name, type, size, and a
    rough char count — so the UI can show the user the file list before they
    commit. Reads file text (bounded) to report chars but writes nothing."""
    docs = discover(roots)
    if limit:
        docs = docs[:limit]
    out = []
    for fp in docs:
        try:
            size = fp.stat().st_size
        except OSError:
            size = 0
        text = extract_text(fp)
        out.append({
            "name": fp.name, "path": str(fp), "ext": fp.suffix.lower(),
            "bytes": size, "chars": len(text), "readable": bool(text),
        })
    return out


def ingest(roots: list[Path] | None = None, store=None,
           max_docs: int | None = None) -> dict:
    """Read the user's documents and mine each for structured facts.

    Idempotent (a file is skipped while its mtime+size are unchanged), best-effort
    (an unreadable file is skipped, never fatal), and reversible (everything is
    tagged source='documents.scan'). Returns per-run counts.
    """
    cfg = settings.documents
    if not cfg.enabled:
        return {"ok": False, "error": "document ingestion disabled (QUILL_DOCUMENTS=0)"}

    from app.events import Event, Modality
    from app.services import confidence as _conf
    from app.services.extractor import extractor
    from app.storage import get_store

    store = store or get_store()
    # Only index facts into the shared vector store when we're using it — a
    # throwaway store (tests) would poison the live index with mismatched ids.
    try:
        live = store is get_store()
    except Exception:
        live = False
    now = time.time()

    docs = discover(roots, max_docs=max_docs)
    state = _load_state()
    ledger = dict(state.get("files") or {})

    counts = {"documents": 0, "chunks": 0, "facts": 0, "skipped": 0,
              "unreadable": 0, "errors": 0}

    for fp in docs:
        key = _file_key(fp)
        sig = _sig(fp)
        if sig and ledger.get(key) == sig:
            counts["skipped"] += 1
            continue

        text = extract_text(fp)
        if not text:
            counts["unreadable"] += 1
            ledger[key] = sig            # remember so we don't retry every run
            continue

        try:
            # One OBSERVED event holds the document (text read verbatim off disk —
            # perfect capture, no model), and every fact anchors back to it.
            ev = Event(
                time=now, modality=Modality.DOCUMENT, raw=text,
                summary=f"[document] {fp.name} ({len(text)} chars)", source=SOURCE,
                meta={"section": "documents", "path": str(fp), "title": fp.name,
                      "ext": fp.suffix.lower()})
            _conf.attach(ev, _conf.OBSERVED, capture=1.0)
            anchor = store.insert(ev)

            doc_facts = 0
            for ch in chunk_text(text):
                try:
                    facts = extractor._extract_text(ch)
                except Exception as exc:
                    print(f"[documents] extract error on {fp.name} ({exc}).")
                    counts["errors"] += 1
                    continue
                doc_facts += _persist_facts(store, facts, anchor, ch, now, index=live)
                counts["chunks"] += 1

            # The document event is itself fully processed — mark it extracted so
            # no general pass re-reads it (the audio extractor already filters by
            # modality, but this keeps the 'unextracted' view honest).
            try:
                store.mark_extracted([anchor], now)
            except Exception:
                pass

            counts["documents"] += 1
            counts["facts"] += doc_facts
            ledger[key] = sig
        except Exception as exc:
            print(f"[documents] ingest error on {fp.name} ({exc}).")
            counts["errors"] += 1

    state["files"] = ledger
    state["last_run"] = now
    _save_state(state)
    print(f"[documents] ingest: {counts}")
    return {"ok": True, **counts}
